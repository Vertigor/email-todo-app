"""
Todo API 路由
"""
from typing import Optional
import uuid
import io
import csv
import json
from urllib.parse import quote
from docx import Document
from datetime import datetime
from calendar import monthrange

from fastapi import APIRouter, HTTPException
from fastapi.responses import Response

from deploy_starter.database import Database
from deploy_starter.models import TodoItem
from deploy_starter.schemas import CompleteTodoRequest, UpdateTodoRequest, CreateTodoRequest
from deploy_starter.utils import get_db_path

router = APIRouter(prefix="/api/todos", tags=["todos"])
db = Database(get_db_path())


@router.get("")
async def get_todos(completed: Optional[bool] = None):
    """获取待办列表（不包含已删除）"""
    todos = db.get_todos(completed, deleted=False)
    return {"todos": [todo.dict() for todo in todos]}


@router.post("")
async def create_todo(request: CreateTodoRequest):
    """手动创建待办事项（无需来源邮件）"""
    due_date = None
    if request.due_date:
        try:
            due_date = datetime.fromisoformat(request.due_date.replace("Z", "+00:00"))
            if due_date.tzinfo:
                due_date = due_date.replace(tzinfo=None)
        except Exception:
            due_date = None

    todo = TodoItem(
        id=str(uuid.uuid4()),
        title=request.title,
        description=request.description or "",
        due_date=due_date,
        created_at=datetime.now(),
        source_email_id="manual",
        source_email_subject="手动创建",
        is_manual=True,
    )
    if db.add_todo(todo):
        return {"success": True, "todo": todo.dict()}
    else:
        raise HTTPException(status_code=500, detail="创建失败")


@router.get("/deleted")
async def get_deleted_todos():
    """获取已删除的待办列表（回收站）"""
    todos = db.get_todos(deleted=True)
    return {"todos": [todo.dict() for todo in todos]}


@router.get("/calendar")
async def get_todos_calendar(year: int, month: int):
    """获取日历视图的待办事项"""
    # 计算月份的第一天和最后一天
    _, last_day = monthrange(year, month)
    start_date = datetime(year, month, 1)
    end_date = datetime(year, month, last_day, 23, 59, 59)
    
    todos = db.get_todos_by_date_range(start_date, end_date)
    
    # 按日期分组
    calendar = {}
    for todo in todos:
        if todo.due_date:
            date_key = todo.due_date.strftime("%Y-%m-%d")
            if date_key not in calendar:
                calendar[date_key] = []
            calendar[date_key].append(todo.dict())
    
    return {"calendar": calendar}


@router.get("/export")
async def export_todos(
    start_date: str,
    end_date: str,
    format: str = "csv",
    completed: Optional[bool] = None,
    count_only: bool = False,
):
    """
    按「创建日期」范围导出待办（用于日报 / 周报编制）。

    - start_date / end_date: YYYY-MM-DD，闭区间
    - format: csv（默认，带 UTF-8 BOM，Excel 可直接打开）或 json
    - completed: 完成状态筛选，None 表示全部
    - count_only: 仅返回该区间待办数量，不返回文件
    """
    try:
        sd = datetime.strptime(start_date, "%Y-%m-%d")
        ed = datetime.strptime(end_date, "%Y-%m-%d")
    except ValueError:
        raise HTTPException(status_code=400, detail="日期格式应为 YYYY-MM-DD")

    if sd > ed:
        raise HTTPException(status_code=400, detail="开始日期不能晚于结束日期")

    try:
        todos = db.get_todos_by_created_range(sd, ed, completed=completed)

        if count_only:
            return {"count": len(todos), "start_date": start_date, "end_date": end_date}

        # 统一把待办转成可导出字段（中文状态 / 格式化时间）
        def _fmt(dt):
            return dt.strftime("%Y-%m-%d %H:%M:%S") if dt else ""

        def _to_row(todo):
            return {
                "title": todo.title or "",
                "status": "已完成" if todo.completed else "待办中",
                "description": todo.description or "",
                "source": "手动录入" if todo.is_manual else "邮件",
                "source_email_subject": todo.source_email_subject or "",
                "source_email_from": todo.source_email_from or "",
                "source_email_to": todo.source_email_to or "",
                "source_email_cc": todo.source_email_cc or "",
                "source_email_date": _fmt(todo.source_email_date),
                "due_date": _fmt(todo.due_date),
                "created_at": _fmt(todo.created_at),
                "completed_at": _fmt(todo.completed_at),
            }

        rows = [_to_row(t) for t in todos]
        # 中文表头（给用户看）；keys 与 headers 一一对应，用于按列取英文键字典的值
        headers = [
            "标题", "状态", "详情", "来源", "邮件主题", "发信人",
            "收件人", "抄送", "邮件时间", "截止时间", "创建时间", "完成时间",
        ]
        keys = [
            "title", "status", "description", "source", "source_email_subject",
            "source_email_from", "source_email_to", "source_email_cc",
            "source_email_date", "due_date", "created_at", "completed_at",
        ]

        if format == "json":
            content = json.dumps(rows, ensure_ascii=False, indent=2)
            return Response(
                content=content.encode("utf-8"),
                media_type="application/json; charset=utf-8",
                headers={
                    "Content-Disposition": f"attachment; filename=todos_{start_date}_{end_date}.json"
                },
            )

        # 默认 CSV：用 utf-8-sig 自动加 BOM，Excel 打开中文不乱码
        buf = io.StringIO()
        writer = csv.writer(buf)
        writer.writerow(headers)
        for r in rows:
            writer.writerow([r[k] for k in keys])
        csv_bytes = buf.getvalue().encode("utf-8-sig")

        return Response(
            content=csv_bytes,
            media_type="text/csv; charset=utf-8",
            headers={
                "Content-Disposition": f"attachment; filename=todos_{start_date}_{end_date}.csv"
            },
        )
    except Exception as exc:
        import traceback
        err = f"导出失败: {type(exc).__name__}: {exc}\n{traceback.format_exc()}"
        print(err, flush=True)
        return Response(
            content=err.encode("utf-8", "replace"),
            status_code=500,
            media_type="text/plain; charset=utf-8",
        )


@router.get("/report")
async def generate_report(
    start_date: str,
    end_date: str,
    period: str = "weekly",
    completed: Optional[bool] = None,
):
    """
    按「创建日期」范围生成工作报告（Word .docx 下载），用于日报 / 周报 / 月报。

    - start_date / end_date: YYYY-MM-DD，闭区间
    - period: daily / weekly / monthly（仅用于标题中文案）
    - completed: 完成状态筛选，None 表示全部
    """
    try:
        sd = datetime.strptime(start_date, "%Y-%m-%d")
        ed = datetime.strptime(end_date, "%Y-%m-%d")
    except ValueError:
        raise HTTPException(status_code=400, detail="日期格式应为 YYYY-MM-DD")

    if sd > ed:
        raise HTTPException(status_code=400, detail="开始日期不能晚于结束日期")

    todos = db.get_todos_by_created_range(sd, ed, completed=completed)
    period_label = {"daily": "日报", "weekly": "周报", "monthly": "月报"}.get(period, "工作报告")
    return _build_docx_report(todos, start_date, end_date, period_label)


def _build_docx_report(todos, start_date: str, end_date: str, period_label: str):
    """把待办列表归纳成一份 Word 工作报告并返回 FastAPI Response。"""
    doc = Document()
    doc.add_heading(f"{period_label}（{start_date} ~ {end_date}）", level=0)

    done = [t for t in todos if t.completed]
    ongoing = [t for t in todos if not t.completed]
    total = len(todos)
    rate = (len(done) / total * 100) if total else 0
    manual = sum(1 for t in todos if t.is_manual)
    mail = total - manual

    doc.add_heading("一、概览", level=1)
    doc.add_paragraph(f"待办总数：{total} 项")
    doc.add_paragraph(f"已完成：{len(done)} 项（完成率 {rate:.0f}%）")
    doc.add_paragraph(f"进行中：{len(ongoing)} 项")
    doc.add_paragraph(f"来源分布：手动录入 {manual} 项，邮件提取 {mail} 项")

    def _src(t):
        return "手动录入" if t.is_manual else f"邮件（{t.source_email_from or '未知发信人'}）"

    doc.add_heading("二、本期已完成", level=1)
    if done:
        for t in done:
            p = doc.add_paragraph(style="List Number")
            p.add_run(t.title)
            if t.description:
                doc.add_paragraph(f"    说明：{t.description}")
            ct = t.completed_at.strftime("%Y-%m-%d %H:%M") if t.completed_at else "—"
            doc.add_paragraph(f"    来源：{_src(t)}　　完成时间：{ct}")
    else:
        doc.add_paragraph("（本周期无已完成待办）")

    doc.add_heading("三、进行中（未完成）", level=1)
    if ongoing:
        for t in ongoing:
            p = doc.add_paragraph(style="List Number")
            p.add_run(t.title)
            if t.description:
                doc.add_paragraph(f"    说明：{t.description}")
            crt = t.created_at.strftime("%Y-%m-%d %H:%M") if t.created_at else "—"
            doc.add_paragraph(f"    来源：{_src(t)}　　创建时间：{crt}")
    else:
        doc.add_paragraph("（本周期无进行中待办）")

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    fname = f"工作报告_{start_date}_{end_date}.docx"
    disp = f"attachment; filename=\"report_{start_date}_{end_date}.docx\"; filename*=UTF-8''{quote(fname)}"
    return Response(
        content=buf.getvalue(),
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": disp},
    )


@router.get("/{todo_id}")
async def get_todo(todo_id: str):
    """获取单个待办详情"""
    todo = db.get_todo_by_id(todo_id)
    if todo:
        return {"todo": todo.dict()}
    else:
        raise HTTPException(status_code=404, detail="待办事项不存在")


@router.put("/{todo_id}")
async def update_todo(todo_id: str, request: UpdateTodoRequest):
    """更新待办事项"""
    # 检查待办是否存在
    todo = db.get_todo_by_id(todo_id)
    if not todo:
        raise HTTPException(status_code=404, detail="待办事项不存在")
    
    # 解析截止日期
    due_date = None
    clear_due_date = False
    
    # 检查 request.due_date：空字符串表示清除，有值则解析
    if request.due_date is not None:
        if request.due_date == "":
            clear_due_date = True
        else:
            try:
                due_date = datetime.fromisoformat(request.due_date.replace("Z", "+00:00"))
                if due_date.tzinfo:
                    due_date = due_date.replace(tzinfo=None)
            except:
                pass
    
    success = db.update_todo(
        todo_id,
        title=request.title,
        description=request.description,
        completed=request.completed,
        due_date=due_date,
        clear_due_date=clear_due_date
    )
    
    if success:
        return {"success": True, "message": "更新成功"}
    else:
        raise HTTPException(status_code=500, detail="更新失败")


@router.put("/{todo_id}/complete")
async def complete_todo(todo_id: str, request: CompleteTodoRequest):
    """标记待办为完成/未完成"""
    success = db.update_todo_completed(todo_id, request.completed)
    if success:
        return {"success": True, "message": "更新成功"}
    else:
        raise HTTPException(status_code=404, detail="待办事项不存在")


@router.put("/{todo_id}/delete")
async def delete_todo(todo_id: str):
    """软删除待办事项（移到回收站）"""
    success = db.soft_delete_todo(todo_id)
    if success:
        return {"success": True, "message": "已移到回收站"}
    else:
        raise HTTPException(status_code=404, detail="待办事项不存在")


@router.put("/{todo_id}/restore")
async def restore_todo(todo_id: str):
    """恢复已删除的待办事项"""
    success = db.restore_todo(todo_id)
    if success:
        return {"success": True, "message": "已恢复"}
    else:
        raise HTTPException(status_code=404, detail="待办事项不存在")
